import os
import fitz
import docx
import pandas as pd
import json
from bs4 import BeautifulSoup
from langchain_core.documents import Document


BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))


# ฟังก์ชันอ่านแต่ละประเภทไฟล์
def load_pdf(path):
    doc = fitz.open(path)
    return "".join([page.get_text() for page in doc])


def load_txt(path):
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def load_docx(path):
    doc = docx.Document(path)
    return "\n".join([p.text for p in doc.paragraphs])


def load_csv(path):
    df = pd.read_csv(path)
    return df.to_string(index=False)


def load_json(path):
    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)
    return json.dumps(data, ensure_ascii=False, indent=2)


def load_html(path):
    with open(path, "r", encoding="utf-8") as f:
        soup = BeautifulSoup(f, "html.parser")
    return soup.get_text()


def load_md(path):
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def load_documents(bot_id):

    data_path = os.path.join(BASE_DIR, "data", bot_id)
    print(f"📂 Looking in: {data_path}")

    documents = []

    if not os.path.exists(data_path):
        print(f" No data folder for bot: {bot_id}")
        return documents

    for file in os.listdir(data_path):

        path = os.path.join(data_path, file)

        try:
            if file.endswith(".pdf"):
                text = load_pdf(path)
            elif file.endswith(".docx"):
                text = load_docx(path)
            elif file.endswith(".txt"):
                text = load_txt(path)
            elif file.endswith(".csv"):
                text = load_csv(path)
            elif file.endswith(".json"):
                text = load_json(path)
            elif file.endswith(".html"):
                text = load_html(path)
            elif file.endswith(".md"):
                text = load_md(path)
            else:
                print(f"❌ Unsupported file: {file}")
                continue

            documents.append(
                Document(
                    page_content=text,
                    metadata={"source": file, "bot_id": bot_id}
                )
            )
            print(f"Loaded: {file}")

        except Exception as e:
            print(f"❌ Error loading {file}: {e}")

    print(f"Total docs: {len(documents)}")
    return documents